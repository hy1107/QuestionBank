from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

def export_to_word(questions, output_path, show_answers=False, title='試卷'):
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    for i, q in enumerate(questions, start=1):
        # Question
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {q['question']}")
        run.font.size = Pt(12)

        # Options
        for letter, key in [('A', 'option_a'), ('B', 'option_b'),
                             ('C', 'option_c'), ('D', 'option_d')]:
            op = doc.add_paragraph(style='List Bullet')
            op.paragraph_format.left_indent = Cm(1)
            run = op.add_run(f"({letter}) {q[key]}")
            run.font.size = Pt(11)

        # Answer (optional)
        if show_answers:
            ans_p = doc.add_paragraph()
            run = ans_p.add_run(f"答案：{q['answer']}")
            run.font.size = Pt(10)
            run.bold = True

        doc.add_paragraph()

    doc.save(output_path)
