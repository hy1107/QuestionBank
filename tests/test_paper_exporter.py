import pytest
import os
from exporter.paper_exporter import export_to_word

SAMPLE_QUESTIONS = [
    {
        'id': 1, 'original_no': '1', 'question': '測試題目一',
        'option_a': '選A', 'option_b': '選B',
        'option_c': '選C', 'option_d': '選D', 'answer': 'A',
    },
    {
        'id': 2, 'original_no': '2', 'question': '測試題目二',
        'option_a': '甲', 'option_b': '乙',
        'option_c': '丙', 'option_d': '丁', 'answer': 'C',
    },
]

def test_export_to_word_creates_file(tmp_path):
    output = str(tmp_path / 'exam.docx')
    export_to_word(SAMPLE_QUESTIONS, output, show_answers=False)
    assert os.path.exists(output)

def test_export_to_word_contains_question_text(tmp_path):
    output = str(tmp_path / 'exam.docx')
    export_to_word(SAMPLE_QUESTIONS, output, show_answers=False)
    from docx import Document
    doc = Document(output)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '測試題目一' in full_text
    assert '測試題目二' in full_text

def test_export_to_word_without_answers_hides_answer(tmp_path):
    output = str(tmp_path / 'exam_no_ans.docx')
    export_to_word(SAMPLE_QUESTIONS, output, show_answers=False)
    from docx import Document
    doc = Document(output)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '答案：' not in full_text

def test_export_to_word_with_answers_shows_answer(tmp_path):
    output = str(tmp_path / 'exam_with_ans.docx')
    export_to_word(SAMPLE_QUESTIONS, output, show_answers=True)
    from docx import Document
    doc = Document(output)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    assert '答案：' in full_text
